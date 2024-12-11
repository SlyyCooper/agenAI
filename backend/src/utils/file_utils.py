# This file contains utility functions for file operations and format conversions
# These functions are used to generate and save research reports in various formats

import aiofiles
import urllib
import mistune
from src.services.firebase.storage_utils import upload_file_to_storage
import io
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
import markdown
import tempfile
import re
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

async def write_to_file(filename: str, text: str) -> None:
    """Asynchronously write text to a file in UTF-8 encoding.

    This function is a core utility used by other functions to write content to files.
    It handles UTF-8 encoding and uses aiofiles for asynchronous file operations.

    Args:
        filename (str): The filename to write to.
        text (str): The text to write.
    """
    # Ensure text is a string
    if not isinstance(text, str):
        text = str(text)

    # Convert text to UTF-8, replacing any problematic characters
    text_utf8 = text.encode('utf-8', errors='replace').decode('utf-8')

    async with aiofiles.open(filename, "w", encoding='utf-8') as file:
        await file.write(text_utf8)

async def write_text_to_md(text: str, filename: str, user_id: str) -> str:
    """Save Markdown to Firebase Storage with user-specific path"""
    return await write_to_firebase_storage(
        text, 
        f"{filename}.md",
        'text/markdown',
        user_id
    )

async def write_md_to_pdf(markdown_content: str, filename: str) -> BytesIO:
    """
    Convert markdown content to PDF format using WeasyPrint
    
    Args:
        markdown_content (str): Markdown content to convert
        filename (str): Base filename for temporary files
        
    Returns:
        BytesIO: PDF content as bytes stream
    """
    try:
        # Convert markdown to HTML with extended features
        html_content = markdown.markdown(
            markdown_content,
            extensions=['extra', 'codehilite', 'tables', 'toc']
        )
        
        # Add HTML wrapper with styling
        html_template = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{filename}</title>
            <style>
                @page {{
                    margin: 1in;
                    size: A4;
                }}
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                    line-height: 1.6;
                    padding: 1em;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin-top: 1.5em;
                    margin-bottom: 0.5em;
                }}
                h1 {{ font-size: 2em; }}
                h2 {{ font-size: 1.5em; }}
                h3 {{ font-size: 1.2em; }}
                p {{ margin: 1em 0; }}
                code {{
                    background: #f8f9fa;
                    padding: 0.2em 0.4em;
                    border-radius: 3px;
                    font-size: 0.9em;
                }}
                pre {{
                    background: #f8f9fa;
                    padding: 1em;
                    border-radius: 5px;
                    overflow-x: auto;
                }}
                blockquote {{
                    border-left: 4px solid #e9ecef;
                    margin: 1em 0;
                    padding-left: 1em;
                    color: #495057;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }}
                th, td {{
                    border: 1px solid #dee2e6;
                    padding: 0.5em;
                    text-align: left;
                }}
                th {{
                    background: #f8f9fa;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Configure fonts
        font_config = FontConfiguration()
        
        # Create PDF using WeasyPrint
        pdf_stream = BytesIO()
        HTML(string=html_template).write_pdf(
            pdf_stream,
            font_config=font_config
        )
        
        pdf_stream.seek(0)
        return pdf_stream
        
    except Exception as e:
        raise ValueError(f"Failed to convert markdown to PDF: {str(e)}")

async def write_md_to_word(markdown_content: str, filename: str) -> BytesIO:
    """
    Convert markdown content to DOCX format
    
    Args:
        markdown_content (str): Markdown content to convert
        filename (str): Base filename (used for document title)
        
    Returns:
        BytesIO: DOCX content as bytes stream
    """
    try:
        # Create new document
        doc = Document()
        
        # Add title
        doc.add_heading(filename, 0)
        
        # Convert markdown to HTML for better formatting
        html_content = markdown.markdown(markdown_content)
        
        # Split content into paragraphs
        paragraphs = html_content.split('\n\n')
        
        # Add paragraphs to document
        for para in paragraphs:
            if para.startswith('<h'):
                # Handle headers
                level = int(para[2])  # Get header level from h1, h2, etc.
                text = para[para.find('>')+1:para.find('</')]
                doc.add_heading(text, level)
            else:
                # Handle regular paragraphs
                doc.add_paragraph(para)
        
        # Save to stream
        docx_stream = BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)
        
        return docx_stream
        
    except Exception as e:
        raise ValueError(f"Failed to convert markdown to DOCX: {str(e)}")

async def write_to_firebase_storage(content: str, filename: str, content_type: str, user_id: str):
    """Write content to Firebase Storage with user-specific path"""
    if not user_id:
        raise ValueError("User ID is required for storage operations")
        
    # Create file-like object in memory
    buffer = io.BytesIO(content.encode('utf-8'))
    # Create user-specific path
    file_path = f"reports/{user_id}/{filename}"
    # Upload to Firebase Storage
    return await upload_file_to_storage(buffer, file_path, content_type)

def sanitize_filename(filename: str) -> str:
    """
    Sanitize a filename to be safe for all operating systems
    
    Args:
        filename (str): The filename to sanitize
        
    Returns:
        str: A sanitized filename safe for all operating systems
    """
    # Replace invalid characters with underscores
    invalid_chars = r'[<>:"/\\|?*\x00-\x1f]'
    sanitized = re.sub(invalid_chars, '_', filename)
    
    # Remove leading/trailing spaces and dots
    sanitized = sanitized.strip('. ')
    
    # Ensure the filename isn't empty after sanitization
    if not sanitized:
        sanitized = 'unnamed_file'
        
    return sanitized

# These utility functions work together to provide a comprehensive report generation system:
# 1. The research report is initially generated in Markdown format.
# 2. write_text_to_md saves the Markdown version of the report.
# 3. write_md_to_pdf converts the Markdown to a styled PDF for easy reading and sharing.
# 4. write_md_to_word converts the Markdown to a DOCX file for users who prefer editable documents.
# 
# This multi-format approach ensures that users can access the research report in their preferred format,
# enhancing the usability and accessibility of the generated content.
